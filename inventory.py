import pandas as pd
import tkinter as tk
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
from docx import Document
import requests
from datetime import date
import pyautogui
import time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

file_cur_data = 'so_lieu_hien_tai.xlsx'
file_saved_data = 'so_lieu_lich_su.xlsx'

df = pd.read_excel(file_cur_data)

df1 = pd.read_excel(file_saved_data)

def date_today():
    today = date.today()
    return today.strftime("%d-%m-%Y")

def them_san_pham():
    new_product = entry_new_product.get()
    df = pd.read_excel(file_cur_data)
    if new_product not in df['Tên hàng'].values:
        new_row = pd.DataFrame(columns=df.columns)
        new_row.loc[0] = [None] * len(df.columns)
        new_row.iloc[0, 0] = new_product
        df = pd.concat([df, new_row], ignore_index=True)
        df.to_excel(file_cur_data, index=False)
        messagebox.showinfo('Thông báo', "Thêm sản phẩm mới thành công!")
    else:
        messagebox.showinfo('Thông báo',"Sản phẩm đã tồn tại!")

def nhap_hang():
    ten_hang = product_to_import.get()
    so_luong_nhap = int(entry_import.get())
    #ngay_nhap = date.today().strftime("%d/%m/%Y")
    ngay_nhap = entry_day_import.get()

    df = pd.read_excel(file_cur_data)
    df1 = pd.read_excel(file_saved_data)
    hang = df.loc[df['Tên hàng']==ten_hang]
    if len(hang) > 0:
        df.loc[df['Tên hàng'] == ten_hang, 'Số lượng'] += so_luong_nhap
        df.loc[df['Tên hàng'] == ten_hang, 'Ngày nhập'] = ngay_nhap
        df.to_excel(file_cur_data, index=False)
        messagebox.showinfo('Thông báo', "Nhập hàng thành công!")
        #tao data de ghi vao file so_lieu_lich_su
        data_lichsu = pd.DataFrame([[ten_hang, so_luong_nhap, ngay_nhap, 'Nhập']], columns=['Tên hàng', 'Số lượng', 'Ngày nhập/xuất', 'Ghi chú'])
        df1 = pd.concat([df1, data_lichsu], ignore_index=True)
        df1.to_excel(file_saved_data, index=False)
    else:
        messagebox.showerror('Lỗi', 'Sản phẩm chưa tồn tại')
        
def xuat_hang():
    ten_hang = product_to_export.get()
    so_luong_xuat = int(entry_export.get())
    #ngay_xuat = date.today().strftime("%d/%m/%Y")
    ngay_xuat = entry_day_export.get()

    df = pd.read_excel(file_cur_data)
    df2 = pd.read_excel(file_saved_data)

    hang = df.loc[df['Tên hàng']==ten_hang]

    if len(hang) > 0:
        so_hien_co = int(df.loc[df['Tên hàng'] == ten_hang, 'Số lượng'])
        if so_luong_xuat > so_hien_co:
            messagebox.showerror('Lỗi', 'Không đủ hàng')
        else:
            df.loc[df['Tên hàng'] == ten_hang, 'Số lượng'] = so_hien_co - so_luong_xuat
            df.loc[df['Tên hàng'] == ten_hang, 'Ngày xuất'] = ngay_xuat
            df.to_excel(file_cur_data, index=False)
            messagebox.showinfo('Thông báo', 'Xuất hàng thành công!')
            #tao data de ghi vao file so_lieu_lich_su
            data_lichsu = pd.DataFrame([[ten_hang, so_luong_xuat, ngay_xuat, 'Xuất']], columns=['Tên hàng', 'Số lượng', 'Ngày nhập/xuất', 'Ghi chú'])
            df2 = pd.concat([df2, data_lichsu], ignore_index=True)
            df2.to_excel(file_saved_data, index=False)
    else:
        messagebox.showerror('Lỗi', 'Sản phẩm chưa tồn tại')

def thong_ke_ton_kho():
    pass

def tim_kiem_theo_ngay():
    ngay = day_entry.get()

    df = pd.read_excel(file_saved_data)
    hang = df.loc[df['Ngày nhập/xuất'] == ngay]
    if hang.empty:
        messagebox.showerror('Lỗi', 'Không tìm thấy kết quả theo ngày {}'.format(ngay))
    else:
        result_text_day.delete(1.0, tk.END)
        result_text_day.insert(tk.END, hang.to_string(index=False))

def tim_kiem_theo_ten():
    ten_hang = product_to_search.get()
    df = pd.read_excel(file_saved_data)
    hang = df.loc[df['Tên hàng'] == ten_hang]
    if hang.empty:
        messagebox.showerror('Lỗi', 'Không tìm thấy kết quả theo sản phẩm {}'.format(ten_hang))
    else:
        result_text_name.delete(1.0, tk.END)
        result_text_name.insert(tk.END, hang.to_string(index=False))

def convert_to_word():
    df = pd.read_excel(file_cur_data)
    time = date_today()

    doc = Document()
    doc.add_heading(f"Thống kê hàng tồn trong kho ngày {time}", level=1)

    #Tao bang va them tieu de
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.style = "Table Grid"
    table.cell(0, 0).text = "Tên hàng"
    table.cell(0, 1).text = "Số lượng"

    for index, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Tên hàng"])
        cells[1].text = str(row["Số lượng"])
    doc.save(f"Thống kê hàng tồn trong kho.docx")
    messagebox.showinfo('Thông báo', 'Tạo báo cáo thành công!')

def send_to_email():
    doc_file = "Thống kê hàng tồn trong kho.docx"
    email_from = 'quangdongtan158@gmail.com'
    email_to = 'letrongquang158@gmail.com'
    password = 'Quang@Le#251098'

    msg = MIMEMultipart()
    msg['From'] = email_from
    msg['To'] = email_to
    msg['Subject'] = 'Thống kê hàng tồn kho'

    with open(doc_file, 'rb') as f:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{doc_file}"')
        msg.attach(part)

    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
    server.login(email_from, password)
    server.sendmail(email_from, email_to, msg.as_string())
    server.quit()

    messagebox.showinfo('Thông báo', 'Gửi email thành công')
def send_to_zalo():
    pyautogui.press('win')
    time.sleep(1)
    pyautogui.write('Zalo')
    time.sleep(1)
    pyautogui.press('enter')
    time.sleep(10)

    pyautogui.click(80, 90)
    time.sleep(1)
    pyautogui.write('Z127-c Trang KH')

    pyautogui.press('enter')
    time.sleep(2)


root = tk.Tk()
root.title("Product Inventory")
root.geometry('1000x1530')

#Tạo Frame chứa phần thêm sản phẩm mới vào kho
frame_addnew = tk.Frame(root, bg="gray", width=800)
frame_addnew.pack(side=tk.TOP, padx=10, pady=10)

add_new_label = tk.Label(frame_addnew, text = "Nhập tên sản phẩm mới:")
add_new_label.grid(row = 0, column = 0)

entry_new_product = tk.Entry(frame_addnew)
entry_new_product.grid(row = 0, column = 1)
addnew_button = tk.Button(frame_addnew, text = "Thêm", command=them_san_pham)
addnew_button.grid(row = 4, column = 0)

#Tạo Frame chứa phần nhập thêm sản phẩm
frame_import = tk.Frame(root, bg="green", width=800)
frame_import.pack(side=tk.TOP, padx=10, pady=10)

choose_import_product = tk.Label(frame_import, text="Chọn sản phẩm muốn nhập:")
choose_import_product.grid(row = 0, column= 0)

#Tạo dropdown list các sản phẩm còn tồn trong kho
im = tk.StringVar()
product_to_import = ttk.Combobox(frame_import, width=30, textvariable=im)
product_to_import.grid(row=0, column=1)

#set value for dropdown box
available_products_to_im = []
for i in df['Tên hàng']:
    available_products_to_im.append(i)
product_to_import['values'] = available_products_to_im
product_to_import.current(0)

#Tạo label số lượng sản phẩm muốn import
import_label = tk.Label(frame_import, text="Số lượng nhập")
import_label.grid(row=1, column=0)
entry_import = tk.Entry(frame_import)
entry_import.grid(row = 1, column = 1)

import_day_label = tk.Label(frame_import, text="Ngày nhập")
import_day_label.grid(row=2, column=0)
entry_day_import = tk.Entry(frame_import)
entry_day_import.grid(row = 2, column = 1)

#Tạo button để nhập sản phẩm
import_button = tk.Button(frame_import, text="Nhập", command=nhap_hang)
import_button.grid(row=3, column=0)

#Tạo Frame chứa phần xuất sản phẩm ra khỏi kho
frame_export = tk.Frame(root, bg="red", width=800)
frame_export.pack(side=tk.TOP, padx=10, pady=10)

choose_export_product = tk.Label(frame_export, text="Chọn sản phẩm muốn xuất:")
choose_export_product.grid(row = 0, column= 0)

#Tạo dropdown list các sản phẩm còn tồn trong kho
ex = tk.StringVar()
product_to_export = ttk.Combobox(frame_export, width=30, textvariable=ex)
product_to_export.grid(row=0, column=1)

#set value for dropdown box
available_products_to_ex = []
for i in df['Tên hàng']:
    available_products_to_ex.append(i)
product_to_export['values'] = available_products_to_ex
product_to_export.current(0)

#Tạo label số lượng sản phẩm muốn export
export_label = tk.Label(frame_export, text="Số lượng xuất")
export_label.grid(row=1, column=0)
entry_export = tk.Entry(frame_export)
entry_export.grid(row = 1, column = 1)

export_day_label = tk.Label(frame_export, text="Ngày xuất")
export_day_label.grid(row=2, column=0)
entry_day_export = tk.Entry(frame_export)
entry_day_export.grid(row = 2, column = 1)

#Tạo button để nhập sản phẩm
export_button = tk.Button(frame_export, text="Xuất", command=xuat_hang)
export_button.grid(row=3, column=0)

#Tạo frame xuất dữ liệu hàng tồn kho hiện có ra file báo cáo dạng word và gửi file qua zalo
convert_frame = tk.Frame(root)
convert_frame.pack(side=tk.TOP, padx=10, pady=10)
Button(convert_frame, text="Hiện tồn", command=thong_ke_ton_kho).pack(side=LEFT)
Button(convert_frame, text="Báo cáo", command=convert_to_word).pack(side=LEFT)
Button(convert_frame, text="Gửi email", command=send_to_email).pack(side=LEFT)
Button(convert_frame, text="Gửi zalo", command=send_to_zalo).pack(side=LEFT)

#Tạo frame tìm kiếm sản phẩm theo ngày hoặc theo sản phẩm
search_frame = tk.Frame(root, width=1000)
search_frame.pack(side=tk.TOP, padx=10, pady=10)
day_label = tk.Label(search_frame, text="Nhập ngày")
day_label.grid(row = 0, column = 0)
day_entry = tk.Entry(search_frame)
day_entry.grid(row = 0, column = 1)
day_button = tk.Button(search_frame, text="Tìm theo ngày", command=tim_kiem_theo_ngay)
day_button.grid(row=1, column=0)
#Hien thi ket qua tim kiem
result_text_day = tk.Text(search_frame, height=10, width=50)
result_text_day.grid(row=1, column=1)

search1_frame = tk.Frame(root, width=1200)
search1_frame.pack(side=tk.TOP, padx=10, pady=10)
product_label = tk.Label(search1_frame, text="Nhập tên sản phẩm")
product_label.grid(row = 0, column = 0)
#day_entry = tk.Entry(search_frame)
#day_entry.grid(row = 1, column = 2)
se = tk.StringVar()
product_to_search = ttk.Combobox(search1_frame, width=10, textvariable=se)
product_to_search.grid(row=0, column=1)

#set value for dropdown box
available_products_to_search = df.iloc[:, 0].values
#print(available_products)
product_to_search['values'] = list(available_products_to_search[0:])
#set default value
#product_to_search.current(1)
product_button = tk.Button(search1_frame, text="Tìm theo tên", command=tim_kiem_theo_ten)
product_button.grid(row=1, column=0)
result_text_name = tk.Text(search1_frame, height=10, width=50)
result_text_name.grid(row=1, column=1)

#Tạo frame button save và close phần mềm.
button = tk.Frame(root, bg="blue")
button.pack(side=tk.TOP, padx=10, pady=10)
#Button(button, text="Save", command = save_change).pack(side=LEFT)
Button(button, text="Đóng", command = root.quit).pack(side=LEFT)


root.mainloop()